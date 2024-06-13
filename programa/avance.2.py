import sys
from datetime import datetime
from PySide6.QtCore import QDate
from PySide6.QtWidgets import QApplication, QMainWindow, QHeaderView,QPushButton, QVBoxLayout, QWidget, QDialog, QLineEdit, QFormLayout, QDateEdit, QTextEdit, QTableWidget, QTableWidgetItem, QLabel, QDialogButtonBox, QComboBox, QMessageBox, QFileDialog
from PySide6.QtCore import Qt
from matplotlib.backends.backend_qt5agg import FigureCanvasQTAgg as FigureCanvas
import matplotlib.pyplot as plt
import pandas as pd
import firebase_admin
from firebase_admin import credentials, firestore
from openpyxl import Workbook
from openpyxl.styles import Alignment
from openpyxl.styles import PatternFill
from openpyxl.styles import Alignment, Border, Side
from openpyxl.utils import get_column_letter
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT




# Inicializa Firebase Admin SDK
cred = credentials.Certificate("C:/Users/anghe/OneDrive/Desktop/firebase_credenciales.txt.json")
firebase_admin.initialize_app(cred, {'databaseURL': "https://www.googleapis.com/robot/v1/metadata/x509/firebase-adminsdk-lss3s%40consultorior-37acd.iam.gserviceaccount.com"})
db = firestore.client()

class VentanaSesiones(QMainWindow):
    instance = None  # Instancia única de VentanaSesiones
    def __init__(self):
        super().__init__()
        if  VentanaSesiones.instance is not None:
            VentanaSesiones.instance.close()  # Cerramos la instancia previa si existe
        VentanaSesiones.instance = self
        self.title = 'Ventana de Sesiones'
        self.initUI()

    def initUI(self):
        self.setWindowTitle(self.title)
        self.setGeometry(100, 100, 400, 300)

        
        self.button_exportar_excel = QPushButton("Exportar a Excel")
        self.button_exportar_excel.clicked.connect(self.exportar_a_excel)

        self.rut_busqueda = QLineEdit()
        self.rut_busqueda.setPlaceholderText("Ingrese RUT para buscar")
        self.rut_busqueda.returnPressed.connect(self.buscar_datos)
        
        self.figure1 = plt.Figure()
        self.canvas1 = FigureCanvas(self.figure1)
        self.ax1 = self.figure1.add_subplot(111)

        self.figure2 = plt.Figure()
        self.canvas2 = FigureCanvas(self.figure2)
        self.ax2 = self.figure2.add_subplot(111)

        self.button_abrir_pacientes = QPushButton("Abrir Pacientes")
        self.button_abrir_pacientes.clicked.connect(self.abrir_ventana_pacientes)

        layout = QVBoxLayout()
        layout.addWidget(self.rut_busqueda)
        layout.addWidget(self.canvas1)
        layout.addWidget(self.canvas2)
        layout.addWidget(self.button_exportar_excel)  # Agregar botón de exportar a Excel
        layout.addWidget(self.button_abrir_pacientes)  # Agregar botón para abrir ventana de pacientes

        container = QWidget()
        container.setLayout(layout)
        self.setCentralWidget(container)

        # Configurar tamaño de ventana
        self.resize(1000, 600)

        # Establecer estilos
        self.setStyleSheet("""
    QPushButton {
        background-color: blue;
        border-radius: 30px;  /* Se corrige la sintaxis y se establece el radio de las esquinas */
        color: white;
        width: 30px;
        height: 30px;  
    }
""")

        # Mostrar la ventana de sesiones al iniciar
        self.buscar_datos()

    
    def abrir_ventana_pacientes(self):
        self.ventana_pacientes = VentanaPacientes()
        self.ventana_pacientes.setStyleSheet(self.styleSheet())  # Aplicar el mismo estilo de la ventana de sesiones a la de pacientes
        self.ventana_pacientes.show()

    def buscar_datos(self):
        rut = self.rut_busqueda.text()
        docs = db.collection('sesiones').stream()
        datos = [doc.to_dict() for doc in docs if doc.to_dict().get('rut') == rut]
        self.visualizar_grafico_dolor(datos)
        self.visualizar_grafico_fuerza(datos)

    def visualizar_grafico_dolor(self, datos):
        self.ax1.clear()  # Limpiar el gráfico antes de actualizar

        sesiones = [int(dato.get('sesion', 1)) for dato in datos]
        dolor_reposo = [int(dato.get('dolor_reposo', 0)) for dato in datos]
        dolor_movimiento = [int(dato.get('dolor_movimiento', 0)) for dato in datos]

        if not sesiones:  # Verificar si la lista de sesiones está vacía
            # Mostrar un mensaje de advertencia al usuario
           
            return

        # Graficar barras de dolor en reposo y en movimiento
        self.ax1.bar(sesiones, dolor_reposo, width=0.4, align='center', label='Dolor en Reposo')
        self.ax1.bar([x + 0.4 for x in sesiones], dolor_movimiento, width=0.4, align='center', label='Dolor en Movimiento')

        # Configuraciones del gráfico
        self.ax1.set_xlabel('Número de Sesión')
        self.ax1.set_ylabel('Dolor')
        self.ax1.set_title('Dolor en Reposo y en Movimiento por Sesión')
        self.ax1.set_xticks(sesiones)
        self.ax1.legend()

        # Actualizar el lienzo del gráfico
        self.canvas1.draw()

    def visualizar_grafico_fuerza(self, datos):
        self.ax2.clear()  # Limpiar el gráfico antes de actualizar

        sesiones = [int(dato.get('sesion', 1)) for dato in datos]
        fuerza_muscular = [int(dato.get('fuerza_muscular', 0)) for dato in datos]

        if not sesiones:  # Verificar si la lista de sesiones está vacía
            # Mostrar un mensaje de advertencia al usuario
            QMessageBox.warning(None, "Visualizar Gráfico", "No hay datos disponibles para graficar.")
            return

        # Graficar barras de fuerza muscular
        self.ax2.bar(sesiones, fuerza_muscular, width=0.4, align='center', label='Fuerza Muscular')

        # Configuraciones del gráfico
        self.ax2.set_xlabel('Número de Sesión')
        self.ax2.set_ylabel('Fuerza Muscular')
        self.ax2.set_title('Fuerza Muscular por Sesión')
        self.ax2.set_xticks(sesiones)
        self.ax2.legend()

        # Actualizar el lienzo del gráfico
        self.canvas2.draw()

    # Dentro del método exportar_a_excel de la clase VentanaSesiones

    def exportar_a_excel(self):
        rut = self.rut_busqueda.text()
        docs = db.collection('sesiones').stream()
        datos_sesiones = [doc.to_dict() for doc in docs if doc.to_dict().get('rut') == rut]

        if not datos_sesiones:
            QMessageBox.warning(self, "Exportar a Excel", "No hay datos disponibles para exportar.")
            return

        # Ordenar los datos por el número de sesión
        datos_sesiones_ordenados = sorted(datos_sesiones, key=lambda x: int(x.get('sesion', 0)))

        # Crear un nuevo libro de Excel y seleccionar la hoja activa
        wb = Workbook()
        ws = wb.active

        # Definir el alto de las filas en 60 píxeles
        for i in range(1, len(datos_sesiones_ordenados) + 2):
            ws.row_dimensions[i].height = 60 / 1.3  # Convertir píxeles a unidades de alto de fila en Excel

        # Ajustar el ancho de las columnas
        column_widths = [130, 220, 138, 138, 138]  # Anchos de las columnas en pixeles
        for i, width in enumerate(column_widths):
            ws.column_dimensions[get_column_letter(i + 1)].width = width / 7.5  # Convertir pixeles a unidades de ancho de columna en Excel

        # Agregar texto "RESPONSABLE" en la celda B1
        ws['B1'] = "RESPONSABLE"

        # Agregar encabezados
        encabezados = ["Sesión", "RUT", "Evolución", "Fisioterapia", "Kinesiterapia"]
        ws.append(encabezados)

        # Agregar datos de sesiones
        for dato in datos_sesiones_ordenados:
            sesion = dato.get("sesion", "")
            rut = dato.get("rut", "")
            evolucion = dato.get("evolucion", "")
            fisioterapia = dato.get("fisioterapia", "")
            kinesiterapia = dato.get("kinesiterapia", "")
            ws.append([sesion, rut, evolucion, fisioterapia, kinesiterapia])

        # Ajustar la altura de las filas y el ancho de las columnas
        for row in ws.iter_rows(min_row=2, max_row=len(datos_sesiones_ordenados) + 2, min_col=1, max_col=5):
            for cell in row:
                cell.alignment = Alignment(wrap_text=True, vertical='top')  # Ajustar el texto para que se ajuste dentro de las celdas
                cell.border = Border(left=Side(style='thin'), right=Side(style='thin'), top=Side(style='thin'), bottom=Side(style='thin'))  # Agregar bordes internos y contorno

        # Ajustar márgenes de la hoja
        ws.page_margins.left = 0.5
        ws.page_margins.right = 0.5
        ws.page_margins.top = 0.5
        ws.page_margins.bottom = 0.5

        # Guardar el archivo Excel
        options = QFileDialog.Options()
        file_name, _ = QFileDialog.getSaveFileName(self, "Guardar como", "", "Excel Files (*.xlsx)", options=options)
        if file_name:
            wb.save(file_name)
            QMessageBox.information(self, "Exportar a Excel", f"Los datos se han exportado a '{file_name}' correctamente.")
class VentanaPacientes(QMainWindow):
    instance = None  # Instancia única de VentanaPacientes
    def __init__(self):
        super().__init__()
        if VentanaPacientes.instance is not None:
            VentanaPacientes.instance.close()  # Cerramos la instancia previa si existe
        VentanaPacientes.instance = self
        self.title = 'Ventana de Pacientes'
        self.initUI()
    def initUI(self):
        self.setWindowTitle(self.title)
        self.setGeometry(100, 100, 400, 300)

        self.button_nuevo_paciente = QPushButton("NUEVO PACIENTE")
        self.button_nuevo_paciente.clicked.connect(self.abrir_formulario_paciente)

        self.table_pacientes = QTableWidget()
        self.table_pacientes.setColumnCount(14)
        self.table_pacientes.setHorizontalHeaderLabels(["Nombre", "Edad", "Teléfono", "RUT", "Fecha", "Médico Tratante", "Cirugía", "Previsión", "Dirección", "Correo ", "MT Consulta" , "Sesiones","IFM Sesiones","IFM General"])
        self.actualizar_tabla_pacientes()

        self.button_abrir_sesiones = QPushButton("Abrir Sesiones")
        self.button_abrir_sesiones.clicked.connect(self.abrir_ventana_sesiones)

        self.rut_busqueda = QLineEdit()
        self.rut_busqueda.setPlaceholderText("Buscar por RUT")
        self.rut_busqueda.returnPressed.connect(self.buscar_paciente)

        self.rut_busqueda.textChanged.connect(self.check_rut_input)

        layout = QVBoxLayout()
        layout.addWidget(self.button_nuevo_paciente)
        layout.addWidget(self.rut_busqueda)
        layout.addWidget(self.table_pacientes)
        layout.addWidget(self.button_abrir_sesiones)

        container = QWidget()
        container.setLayout(layout)
        self.setCentralWidget(container)

        self.resize(1230, 600)

        self.setStyleSheet("""
            QPushButton {
                background-color: blue;
                border-radius: 30x;
                color: white;
                width: 30px;
                height: 30px;
            }
        """)
    def check_rut_input(self, text):
            if text.strip() == "1":
                self.rut_busqueda.clear()  # Limpiar el campo de búsqueda
                self.actualizar_tabla_pacientes()  # Actualizar la tabla de pacientes
                
    def abrir_ventana_sesiones(self):
        self.ventana_sesiones = VentanaSesiones()
        self.ventana_sesiones.show()


    def abrir_formulario_paciente(self):
        self.formulario_paciente = FormularioPaciente()
        self.formulario_paciente.show()


    def actualizar_tabla_pacientes(self):
        pacientes = self.recuperar_pacientes()
        self.table_pacientes.setRowCount(len(pacientes))
        for i, paciente in enumerate(pacientes):
            self.table_pacientes.setItem(i, 0, QTableWidgetItem(paciente.get("nombre", "")))
            self.table_pacientes.setItem(i, 1, QTableWidgetItem(paciente.get("edad", "")))
            self.table_pacientes.setItem(i, 2, QTableWidgetItem(paciente.get("telefono", "")))
            self.table_pacientes.setItem(i, 3, QTableWidgetItem(paciente.get("rut", "")))
            self.table_pacientes.setItem(i, 4, QTableWidgetItem(paciente.get("fecha", "")))
            self.table_pacientes.setItem(i, 5, QTableWidgetItem(paciente.get("medicoTratante", "")))
            self.table_pacientes.setItem(i, 6, QTableWidgetItem(paciente.get("cirugia", "")))
            self.table_pacientes.setItem(i, 7, QTableWidgetItem(paciente.get("prevision", "")))
            self.table_pacientes.setItem(i, 8, QTableWidgetItem(paciente.get("direccion", "")))
            self.table_pacientes.setItem(i, 9, QTableWidgetItem(paciente.get("correo", "")))
            self.table_pacientes.setItem(i, 10, QTableWidgetItem(paciente.get("motivo_consulta", "")))
            # Ajustar el ancho de las columnas para que el contenido se muestre completamente
            self.table_pacientes.resizeColumnsToContents()
            # Permitir que el contenido de las celdas se expanda y ajuste automáticamente
            self.table_pacientes.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)

            button_nueva_sesion = QPushButton("Nueva Sesión")
            button_nueva_sesion.clicked.connect(lambda _, rut=paciente.get("rut", ""): self.nueva_sesion(rut))  # Conectar el clic del botón a una función
            self.table_pacientes.setCellWidget(i, 11, button_nueva_sesion)  # Establecer el botón en la celda correspondiente
            # Agregar botón "Informe"
            if self.table_pacientes.cellWidget(i, 13) is None:
                button_informe = QPushButton("IFM general")
                button_informe.clicked.connect(lambda _, rut=paciente.get("rut", ""): self.abrir_formulario_informe(rut))
                self.table_pacientes.setCellWidget(i, 13, button_informe)

            button_generar_informe = QPushButton("IFM sesiones")
            button_generar_informe.clicked.connect(lambda _, rut=paciente.get("rut", ""): self.generar_informe(rut))  # Conectar el clic del botón a una función
            self.table_pacientes.setCellWidget(i, 12, button_generar_informe)  # Establecer el botón en la celda correspondiente
    

    def nueva_sesion(self, rut):
        # Esta función maneja lo que sucede cuando se hace clic en el botón "Nueva Sesión"
        self.formulario_sesion = FormularioSesion()
        self.formulario_sesion.rut.setText(rut)  # Establecer el rut del paciente en el formulario de sesión
        self.formulario_sesion.show()

    def generar_informe(self, rut):
        docs = db.collection('sesiones').stream()
        datos = [doc.to_dict() for doc in docs if doc.to_dict().get('rut') == rut]

        if not datos:
            QMessageBox.warning(self, "Generar Informe", "No hay datos disponibles para generar el informe.")
            return

        # Ordenar los datos por número de sesión
        datos.sort(key=lambda x: x.get('sesion', ''))

        paciente_doc = db.collection('pacientes').document(rut).get()
        paciente = paciente_doc.to_dict() if paciente_doc.exists else None

        doc = Document()

        if paciente:
            doc.add_heading(f"Informe del Paciente: {paciente.get('nombre', '')}", level=1)
            doc.add_paragraph(f"RUT: {paciente.get('rut', '')}")
            doc.add_paragraph(f"Edad: {paciente.get('edad', '')}")
            doc.add_paragraph(f"Teléfono: {paciente.get('telefono', '')}")
            doc.add_paragraph(f"Fecha: {paciente.get('fecha', '')}")
            doc.add_paragraph(f"Médico Tratante: {paciente.get('medicoTratante', '')}")
            doc.add_paragraph(f"Cirugía: {paciente.get('cirugia', '')}")
            doc.add_paragraph(f"Previsión: {paciente.get('prevision', '')}")
            doc.add_paragraph(f"Dirección: {paciente.get('direccion', '')}")
            doc.add_paragraph(f"Correo Electrónico: {paciente.get('correo', '')}")
            doc.add_paragraph(f"Motivo de Consulta: {paciente.get('motivo_consulta', '')}")

        doc.add_heading("Sesiones", level=2)
        for dato in datos:
            doc.add_heading(f"Sesión {dato.get('sesion', '')}", level=3)
            doc.add_paragraph(f"Fecha: {dato.get('fecha', '')}")
            doc.add_paragraph(f"Observaciones: {dato.get('observaciones', '')}")
            doc.add_paragraph(f"Evolución: {dato.get('evolucion', '')}")
            doc.add_paragraph(f"Examen Físico: {dato.get('examen_fisico', '')}")
            doc.add_paragraph(f"Evaluación Muscular: {dato.get('evaluacion_muscular', '')}")
            doc.add_paragraph(f"Medición Articular: {dato.get('medicion_articular', '')}")
            doc.add_paragraph(f"Fuerza Muscular: {dato.get('fuerza_muscular', '')}")
            doc.add_paragraph(f"Dolor en Reposo: {dato.get('dolor_reposo', '')}")
            doc.add_paragraph(f"Dolor en Movimiento: {dato.get('dolor_movimiento', '')}")
            doc.add_paragraph(f"Actitud Funcional: {dato.get('actitud_funcional', '')}")

        options = QFileDialog.Options()
        file_name, _ = QFileDialog.getSaveFileName(self, "Guardar Informe como", "", "Word Files (*.docx)", options=options)
        if file_name:
            doc.save(file_name)
            QMessageBox.information(self, "Generar Informe", f"El informe se ha generado y guardado en '{file_name}' correctamente.")



    def abrir_formulario_informe(self, rut):
        # Recuperar los datos del paciente según el RUT
        paciente = self.recuperar_datos_paciente(rut)
        if paciente:
            # Obtener los datos del paciente
            nombre = paciente.get("nombre", "")
            edad = paciente.get("edad", "")
            motivo_consulta = paciente.get("motivo_consulta", "")
            enfermedad_actual = paciente.get("enfermedad_actual", "")
            examen_fisico = paciente.get("examen_fisico", "")  

            # Crear una instancia de GenerarInformeFormulario
            self.formulario_informe = GenerarInformeFormulario()

            # Establecer los datos del paciente en la ventana de GenerarInformeFormulario
            self.formulario_informe.set_datos(nombre, rut, edad, examen_fisico, enfermedad_actual, motivo_consulta)

            # Mostrar la ventana de GenerarInformeFormulario
            self.formulario_informe.show()
        else:
            QMessageBox.warning(self, "Error", "No se encontró ningún paciente con el RUT proporcionado.")


    def recuperar_datos_paciente(self, rut):
        # Función para recuperar los datos del paciente según el RUT desde la base de datos

        # Inicializamos las variables para almacenar los datos recuperados
        datos_paciente = {}

        # Consultamos la colección de pacientes para obtener los datos básicos del paciente
        pacientes_ref = db.collection("pacientes").where("rut", "==", rut).limit(1).stream()
        for doc in pacientes_ref:
            datos_paciente = doc.to_dict()

        # Consultamos la colección de sesiones para obtener el examen físico más reciente del paciente
        examen_fisico = ""
        sesiones_ref = db.collection("sesiones").where("rut", "==", rut).order_by("fecha", direction=firestore.Query.DESCENDING).limit(1).stream()
        for doc in sesiones_ref:
            datos_sesion = doc.to_dict()
            examen_fisico = datos_sesion.get("examen_fisico", "")

        # Retornamos los datos del paciente junto con el examen físico de la última sesión
        datos_paciente["examen_fisico"] = examen_fisico
        return datos_paciente



    def recuperar_pacientes(self):
            pacientes = []
            docs = db.collection("pacientes").stream()
            for doc in docs:
                paciente = doc.to_dict()
                # Asegurarse de que la clave 'sesiones' exista en el diccionario del paciente
                paciente.setdefault('sesiones', [])
                pacientes.append(paciente)
            return pacientes

    def buscar_paciente(self):
            rut = self.rut_busqueda.text()
            rows = self.table_pacientes.rowCount()
            for i in range(rows):
                item = self.table_pacientes.item(i, 3)
                if item and item.text() == rut:
                    if i != 0:
                        self.mover_fila_al_principio(i)
                    return
            QMessageBox.warning(self, "Buscar Paciente", "No se encontró ningún paciente con el RUT proporcionado.")

    def mover_fila_al_principio(self, row):
        column_count = self.table_pacientes.columnCount()
        new_row = []
        for col in range(column_count):
            item = self.table_pacientes.takeItem(row, col)
            if item:
                new_row.append(item)
        self.table_pacientes.insertRow(0)
        for col, item in enumerate(new_row):
            self.table_pacientes.setItem(0, col, item)
        # Restablecer los widgets de celda para las columnas de los botones
        self.restablecer_botones_celda(0)
        self.table_pacientes.removeRow(row)

    def restablecer_botones_celda(self, row):
        # Verificar si hay un botón "Nueva Sesión" en la celda
        if self.table_pacientes.cellWidget(row, 11) is None:
            button_nueva_sesion = QPushButton("Nueva Sesión")
            button_nueva_sesion.clicked.connect(lambda _, rut=self.table_pacientes.item(row, 3).text(): self.nueva_sesion(rut))
            self.table_pacientes.setCellWidget(row, 11, button_nueva_sesion)
        # Verificar si hay un botón "IFM sesiones" en la celda
        if self.table_pacientes.cellWidget(row, 12) is None:
            button_generar_informe = QPushButton("IFM sesiones")
            button_generar_informe.clicked.connect(lambda _, rut=self.table_pacientes.item(row, 3).text(): self.generar_informe(rut))
            self.table_pacientes.setCellWidget(row, 12, button_generar_informe)
        # Verificar si hay un botón "IFM general" en la celda
        if self.table_pacientes.cellWidget(row, 13) is None:
            button_informe = QPushButton("IFM general")
            button_informe.clicked.connect(lambda _, rut=self.table_pacientes.item(row, 3).text(): self.abrir_formulario_informe(rut))
            self.table_pacientes.setCellWidget(row, 13, button_informe)



class FormularioPaciente(QDialog):
    def __init__(self):
        super(FormularioPaciente, self).__init__()
        self.setWindowTitle("Nuevo Paciente")

        self.nombre = QLineEdit()
        self.edad = QLineEdit()
        self.telefono = QLineEdit()
        self.rut = QLineEdit()
        self.fechaNacimiento = QDateEdit()
        self.medicoTratante = QLineEdit()
        self.cirugia = QLineEdit()
        self.prevision = QLineEdit()
        self.direccion = QLineEdit()
        self.correo = QLineEdit()
        self.enfermedad_actual  = QTextEdit()
        self.motivo_consulta = QTextEdit()

        layout = QFormLayout(self)
        layout.addRow(QLabel("Formulario de Paciente"))
        layout.addRow("Nombre", self.nombre)
        layout.addRow("Edad", self.edad)
        layout.addRow("Teléfono", self.telefono)
        layout.addRow("RUT", self.rut)
        layout.addRow("Fecha de Nacimiento", self.fechaNacimiento)
        layout.addRow("Médico Tratante", self.medicoTratante)
        layout.addRow("Cirugía", self.cirugia)
        layout.addRow("Previsión", self.prevision)
        layout.addRow("Dirección", self.direccion)
        layout.addRow("Correo Electrónico", self.correo)
        layout.addRow("enferdad actual", self.enfermedad_actual)
        layout.addRow("motivo de consulta", self.motivo_consulta)

        buttonBox = QDialogButtonBox(QDialogButtonBox.Save | QDialogButtonBox.Cancel)
        buttonBox.accepted.connect(self.guardar_datos)
        buttonBox.rejected.connect(self.reject)
        layout.addRow(buttonBox)

    def guardar_datos(self):
        data = {
            "nombre": self.nombre.text(),
            "edad": self.edad.text(),
            "telefono": self.telefono.text(),
            "rut": self.rut.text(),
            "fecha": self.fechaNacimiento.date().toString(Qt.ISODate),
            "medicoTratante": self.medicoTratante.text(),
            "cirugia": self.cirugia.text(),
            "prevision": self.prevision.text(),
            "direccion": self.direccion.text(),
            "correo": self.correo.text(),
            "enfermedad_actual": self.enfermedad_actual.toPlainText(),
            "motivo_consulta": self.motivo_consulta.toPlainText()
            
        }

        try:
            db.collection("pacientes").add(data)
            QMessageBox.information(self, "Guardar Datos", "Los datos del paciente se han guardado correctamente.")
            self.accept()
        except Exception as e:
            QMessageBox.critical(self, "Error", f"No se pudieron guardar los datos del paciente: {str(e)}")

class FormularioSesion(QDialog):
    def __init__(self):
        super(FormularioSesion, self).__init__()
        self.setWindowTitle("Nueva Sesión")

       
        self.rut = QLineEdit()
        self.Ncesion = QLineEdit()
        self.fecha = QDateEdit()
        self.fecha.setDate(QDate.currentDate()) 
        self.medico = QLineEdit()
        self.observaciones = QTextEdit()
        self.evolucion = QTextEdit()
        self.examen_fisico = QTextEdit()
        self.evaluacion_muscular = QComboBox()
        self.evaluacion_muscular.addItems(["Normal", "Anormal"])
        self.medicion_articular = QLineEdit()
        self.fuerza_muscular = QComboBox()
        self.fuerza_muscular.addItems(["0", "1", "2", "3", "4", "5"])
        self.dolor_reposo = QComboBox()
        self.dolor_reposo.addItems(["0", "1", "2", "3", "4", "5", "6", "7", "8", "9", "10"])
        self.dolor_movimiento = QComboBox()
        self.dolor_movimiento.addItems(["0", "1", "2", "3", "4", "5", "6", "7", "8", "9", "10"])
        self.actitud_funcional = QComboBox()
        self.actitud_funcional.addItems(["Normal", "Anormal"])

        layout = QFormLayout(self)
        layout.addRow(QLabel("Formulario de Sesión"))
        layout.addRow("RUT", self.rut)
        layout.addRow("N° de Sesión", self.Ncesion)
        layout.addRow("Fecha", self.fecha)
        layout.addRow("Médico", self.medico)
        layout.addRow("Observaciones", self.observaciones)
        layout.addRow("Evolución", self.evolucion)
        layout.addRow("Examen Físico", self.examen_fisico)
        layout.addRow("Evaluación Muscular", self.evaluacion_muscular)
        layout.addRow("Medición Articular", self.medicion_articular)
        layout.addRow("Fuerza Muscular", self.fuerza_muscular)
        layout.addRow("Dolor en Reposo", self.dolor_reposo)
        layout.addRow("Dolor en Movimiento", self.dolor_movimiento)
        layout.addRow("Actitud Funcional", self.actitud_funcional)

        buttonBox = QDialogButtonBox(QDialogButtonBox.Save | QDialogButtonBox.Cancel)
        buttonBox.accepted.connect(self.guardar_datos)
        buttonBox.rejected.connect(self.reject)
        layout.addRow(buttonBox)

    def guardar_datos(self):
        data = {
            "rut": self.rut.text(),
            "sesion": self.Ncesion.text(),
            "fecha": self.fecha.date().toString(Qt.ISODate),
            "medico": self.medico.text(),
            "observaciones": self.observaciones.toPlainText(),
            "evolucion": self.evolucion.toPlainText(),
            "examen_fisico": self.examen_fisico.toPlainText(),
            "evaluacion_muscular": self.evaluacion_muscular.currentText(),
            "medicion_articular": self.medicion_articular.text(),
            "fuerza_muscular": self.fuerza_muscular.currentText(),
            "dolor_reposo": self.dolor_reposo.currentText(),
            "dolor_movimiento": self.dolor_movimiento.currentText(),
            "actitud_funcional": self.actitud_funcional.currentText()
        }

        try:
            db.collection("sesiones").add(data)
            QMessageBox.information(self, "Guardar Datos", "Los datos de la sesión se han guardado correctamente.")
            self.accept()
        except Exception as e:
            QMessageBox.critical(self, "Error", f"No se pudieron guardar los datos de la sesión: {str(e)}")
class GenerarInformeFormulario(QWidget):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Generar Informe Médico")
        self.initUI()

    def initUI(self):
        layout = QFormLayout()

        self.nombre_input = QLineEdit()
        layout.addRow("Nombre:", self.nombre_input)

        self.rut_input = QLineEdit()
        layout.addRow("RUT:", self.rut_input)

        self.edad_input = QLineEdit()
        layout.addRow("EDAD:", self.edad_input)

        self.examen_fisico_input =  QTextEdit()
        layout.addRow("Examen Fisico:", self.examen_fisico_input)

        self.enfermedad_actual_input = QTextEdit()
        layout.addRow("Enfermedad actual :", self.enfermedad_actual_input)

        self.motivo_consulta_input = QTextEdit()
        layout.addRow("Motivo de Consulta:", self.motivo_consulta_input)

        self.diagnostico_input = QTextEdit()
        layout.addRow("Diagnóstico:", self.diagnostico_input)

        self.indicaciones_input = QTextEdit()
        layout.addRow("Indicaciones:", self.indicaciones_input)

        self.generar_informe_button = QPushButton("Generar Informe")
        self.generar_informe_button.clicked.connect(self.generar_informe)
        layout.addWidget(self.generar_informe_button)

        self.setLayout(layout)  # Establecer el diseño en la ventana

    def set_datos(self, nombre, rut, edad, examen_fisico, enfermedad_actual, motivo_consulta):
        # Método para establecer los datos del paciente en la ventana
        self.nombre_input.setText(nombre)
        self.rut_input.setText(rut)
        self.edad_input.setText(edad)
        self.examen_fisico_input.setPlainText(examen_fisico)
        self.enfermedad_actual_input.setPlainText(enfermedad_actual)
        self.motivo_consulta_input.setPlainText(motivo_consulta)
    def generar_informe(self):
        # Obtener los valores de los campos de diagnóstico e indicaciones
        diagnostico = self.diagnostico_input.toPlainText()
        indicaciones = self.indicaciones_input.toPlainText()

        # Verificar si los campos requeridos están vacíos
        if not diagnostico or not indicaciones:
            QMessageBox.warning(self, "Campos Incompletos", "Por favor complete los campos de diagnóstico e indicaciones.")
            return

        # Obtener los valores de los otros campos
        nombre = self.nombre_input.text()
        rut = self.rut_input.text()
        edad = self.edad_input.text()
        examen_fisico = self.examen_fisico_input.toPlainText()
        enfermedad_actual = self.enfermedad_actual_input.toPlainText()
        motivo_consulta = self.motivo_consulta_input.toPlainText()

        # Crear el documento de Word
        doc = Document()

        # Agregar encabezado con la fecha actual
        fecha_actual = datetime.now().strftime("%d/%m/%Y")
        doc.add_heading(f"Fecha: {fecha_actual}", level=1)

        # Agregar información del paciente en la esquina superior izquierda
        doc.add_paragraph(f"Nombre: {nombre}")
        doc.add_paragraph(f"RUT: {rut}")
        doc.add_paragraph(f"Edad: {edad}")

        # Agregar título "Informe Médico"
        title_paragraph = doc.add_heading("INFORME MÉDICO", level=1)
        title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Agregar un nuevo run al párrafo del título
        run = title_paragraph.add_run()

        # Establecer el tamaño de fuente del run
        run.font.size = Pt(36)

        # Agregar sección de motivo de consulta
        doc.add_heading("Motivo de Consulta", level=2)
        doc.add_paragraph(motivo_consulta)

        # Agregar sección de enfermedad actual
        doc.add_heading("Enfermedad Actual", level=2)
        doc.add_paragraph(enfermedad_actual)

        # Agregar sección de examen físico
        doc.add_heading("Examen Físico", level=2)
        doc.add_paragraph(examen_fisico)

        # Agregar sección de diagnóstico
        doc.add_heading("Diagnóstico", level=2)
        doc.add_paragraph(diagnostico)

        # Agregar sección de indicaciones
        doc.add_heading("Indicaciones", level=2)
        doc.add_paragraph(indicaciones)

        # Agregar firma del médico centrada
        doc.add_paragraph("\n\n\n", style='BodyText')
        doc.add_paragraph("Dr. [Nombre del Médico]", style='BodyText').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_paragraph("Traumatología y Ortopedia", style='BodyText').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Guardar el documento como archivo .docx
        options = QFileDialog.Options()
        file_name, _ = QFileDialog.getSaveFileName(self, "Guardar Informe como", "", "Word Files (*.docx)", options=options)
        if file_name:
            doc.save(file_name)
            QMessageBox.information(self, "Generar Informe", f"El informe se ha generado y guardado en '{file_name}' correctamente.")

if __name__ == "__main__":
    app = QApplication([])
    window = VentanaSesiones()
    window.show()
    sys.exit(app.exec())
